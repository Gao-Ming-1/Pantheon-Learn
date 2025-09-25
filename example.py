import os
from docx import Document
from openai import OpenAI
import time

# 用你的 Unlimited API Key
client = OpenAI(
    api_key="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",
    base_url="https://openrouter.ai/api/v1"
)

# 各科目的示例题
SUBJECT_EXAMPLES = {
    "Pure Chemistry": """
Example O-Level Pure Chemistry MCQs:
25.00 cm³ of 0.100 mol·dm⁻³ HCl is titrated against NaOH. The HCl is completely neutralised by 30.00 cm³ of NaOH. What is the concentration of the NaOH?
A. 0.067 mol·dm⁻³
B. 0.083 mol·dm⁻³
C. 0.100 mol·dm⁻³
D. 0.120 mol·dm⁻³

In paper chromatography the solvent front moves 9.0 cm from the origin and a dye spot moves 3.6 cm from the origin. What is the Rf value of the dye?
A. 0.25
B. 0.40
C. 0.60
D. 2.50

A mixture of ethyl acetate and water is to be separated. Which method is most suitable and which physical test best confirms the purity of the separated ethyl acetate?
A. Distillation; measure boiling point.
B. Separating funnel; measure boiling point.
C. Filtration; check refractive index.
D. Separating funnel; check for immiscibility with water.

A gas denser than air is released into an open container. After a long time at room temperature, which is most likely?
A. The dense gas will remain concentrated at the bottom.
B. The dense gas will mix uniformly with the air by diffusion.
C. The dense gas will rise and collect at the top.
D. The dense gas will react with oxygen and disappear.

A compound contains 2 carbon atoms (only ¹²C), 1 hydrogen atom (¹H only) and two chlorine atoms which may be ³⁵Cl or ³⁷Cl. How many distinct relative molecular masses are possible?
A. 2
B. 3
C. 4
D. 5

Which explains why ionic lattice A has a higher melting point than ionic lattice B?
A. Ions in A are larger than ions in B.
B. A has ions with higher charges than B.
C. A is covalent while B is ionic.
D. A has free electrons while B does not.

The outer-shell electrons diagram of molecule X shows one central atom with six valence electrons bonded to two atoms each contributing one electron pair, and one lone pair on the central atom. Which element could be the central atom?
A. Carbon (proton no. 6)
B. Nitrogen (7)
C. Oxygen (8)
D. Sulfur (16)

A solid substance K conducts electricity in the solid state and reacts with dilute HCl to give a gas that burns with a pop. Which substance could K be?
A. Graphite
B. Calcium carbonate
C. Sodium chloride
D. Zinc

Which general property can you deduce about a metal element M from the fact that metallic bonding predominates in its structure?
A. M is always very reactive.
B. M is a good conductor of heat and electricity.
C. M must have a low melting point.
D. M is always soft.

Silicon dioxide (SiO₂) in powdered form is heated strongly in oxygen. Which set of properties describes SiO₂'s behavior and physical character?
A. Hard solid; stable — no combustion
B. Soft solid; decomposes to silicon and oxygen
C. Hard solid; combusts producing CO₂
D. Low-melting solid; melts to give a liquid residue

Which statement about diamond and graphite is correct?
A. Diamond and graphite are allotropes of carbon with different structures.
B. Graphite does not conduct electricity but diamond does.
C. Both have three-dimensional networks of single C–C bonds only.
D. Diamond contains layers of hexagonal rings similar to graphite.

What is the relative molecular mass (Mr) of butanoic acid, C₄H₈O₂? (C=12, H=1, O=16)
A. 72
B. 88
C. 90
D. 102

Which compound has the highest percentage by mass of oxygen?
A. Fe₂O₃ (Mr 160)
B. H₂O₂ (Mr 34)
C. CO₂ (Mr 44)
D. C₆H₁₂O₆ (Mr 180)

For which gaseous reaction does the total number of moles of gas decrease the most?
A. 2NO₂(g) → N₂O₄(g)
B. N₂(g) + 3H₂(g) → 2NH₃(g)
C. CH₄(g) + 2O₂(g) → CO₂(g) + 2H₂O(g)
D. 2SO₂(g) + O₂(g) → 2SO₃(g)

25.0 cm³ of 0.150 mol·dm⁻³ H₂SO₄ completely neutralises 50.0 cm³ of KOH. What is the concentration of the KOH?
A. 0.0375 mol·dm⁻³
B. 0.0750 mol·dm⁻³
C. 0.150 mol·dm⁻³
D. 0.300 mol·dm⁻³

A metal sample of mass 5.0 g reacts to give an oxide of mass 6.8 g. What is the percentage yield if the theoretical oxide mass is 7.0 g?
A. 68%
B. 97%
C. 97.1%
D. 100%

Which ionic equation represents neutralisation of aqueous NH₃ (ammonia) with aqueous HCl?
A. NH₄⁺ + OH⁻ → NH₃ + H₂O
B. NH₃ + H⁺ → NH₄⁺
C. NH₃ + Cl⁻ → NH₃Cl⁻
D. NH₄Cl → NH₄⁺ + Cl⁻

Which pair of aqueous solutions gives a white precipitate on mixing?
A. Silver nitrate + sodium chloride
B. Sodium sulfate + barium nitrate
C. Potassium nitrate + silver nitrate
D. Sodium hydroxide + ammonium chloride

Which reaction is the best laboratory method to prepare an insoluble metal sulfate precipitate?
A. Mixing a soluble metal chloride with dilute sulfuric acid.
B. Mixing a soluble metal nitrate with dilute sulfuric acid.
C. Mixing a soluble metal carbonate with dilute sulfuric acid.
D. Heating the metal in sulfuric acid.

A solid X gives effervescence with dilute HCl and the gas turns limewater milky; warmed with aqueous sodium hydroxide first produces a gas that turns damp red litmus blue. Which could X be?
A. NaHCO₃
B. NH₄Cl
C. CaCO₃
D. CuCO₃

Which of the following are redox reactions?

Zn + 2HCl → ZnCl₂ + H₂

CuSO₄ + Fe → FeSO₄ + Cu

H₂O₂ → H₂O + ½O₂ (decomposition)
A. 1 and 2 only
B. 1 and 3 only
C. 2 and 3 only
D. 1, 2 and 3

In the electrolysis of molten sodium chloride using inert electrodes, which occurs?
A. Oxygen is liberated at the anode.
B. Sodium metal collects at the cathode.
C. Chloride ions are discharged at the cathode.
D. The electrolyte becomes more acidic.

Which electrochemical cell experiment will show the greatest change in concentration of one ionic species over time (using inert electrodes)?
A. Electrolysis of dilute copper(II) sulfate with carbon electrodes.
B. Electrolysis of aqueous sodium chloride with platinum electrodes.
C. Electrolysis of concentrated sodium chloride with carbon electrodes.
D. Electrolysis of dilute sulfuric acid with platinum electrodes.

Two metal rods touch in dilute sulfuric acid; hydrogen bubbles evolve at one rod only. Which is true?
A. The rod with bubbles is more reactive.
B. Electrons flow from the rod producing hydrogen to the other rod.
C. The rod with bubbles is being plated by the other metal.
D. The rod with bubbles is noble and unreactive.

The element with proton number 20 (calcium) has chemical properties most similar to which proton number?
A. 2
B. 12
C. 19
D. 38

An element P has high melting point and forms coloured ions and acts as a catalyst in some reactions. P is most likely:
A. An alkali metal
B. A halogen
C. A transition metal
D. A noble gas

Chromium(III) oxide is to be reduced to chromium metal. Which reducing agent is most likely to reduce Cr₂O₃?
A. Copper metal
B. Carbon monoxide
C. Zinc metal
D. Lead metal

A hot stream of hydrogen is passed over powdered CuO and the exit gas is bubbled through cold water. A colourless liquid condenses and collects. Which was formed in the U-tube?
A. Water (H₂O)
B. Methanol (CH₃OH)
C. Copper hydroxide solution
D. Carbon dioxide dissolved

The reaction H₂ + I₂ ⇌ 2HI is exothermic. Increasing the temperature shifts the equilibrium:
A. To the right (more HI)
B. To the left (more H₂ and I₂)
C. No change in position
D. Reaction stops completely

Which energy profile corresponds to an exothermic reaction with a single-step mechanism and low activation energy? (select by description)
A. Products lower than reactants; small hump.
B. Products higher than reactants; large hump.
C. Products lower than reactants; large hump.
D. Products higher than reactants; small hump.

For which reaction is changing the pressure least likely to affect the rate at room conditions?
A. CO(g) + ½O₂(g) → CO₂(g)
B. 2NO₂(g) → N₂O₄(g)
C. NaOH(aq) + HCl(aq) → NaCl(aq) + H₂O(l)
D. H₂(g) + Cl₂(g) → 2HCl(g)

The decomposition of H₂O₂ catalysed by MnO₂ is followed by measuring O₂ volume vs time. Three experiments use the same total moles of H₂O₂ but different volumes: (i) 50 cm³ of 2.0 M, (ii) 100 cm³ of 1.0 M, (iii) 100 cm³ of 2.0 M. Which two experiments produce identical initial rates assuming rate depends only on concentration?
A. (i) and (ii)
B. (i) and (iii)
C. (ii) and (iii)
D. None are identical

Octane (C₈H₁₈) undergoes free-radical halogenation with Cl₂ under UV light. The major type of reaction is:
A. Addition across a C=C bond
B. Substitution at a C–H bond
C. Polymerisation to long chains
D. Electrophilic addition to give Cl–C₈H₁₉

Which structure is an isomer of 2-propanol (isopropyl alcohol)? (choose description)
A. CH₃CH(OH)CH₃
B. CH₃CH₂CH₂OH
C. (CH₃)₃C–OH
D. CH₂=CH–OH

Which organic compound reacts with sodium carbonate to release CO₂ gas?
A. Ethanol (CH₃CH₂OH)
B. Ethanoic acid (CH₃COOH)
C. 2-propanone (acetone)
D. Ethyl ethanoate (CH₃COOCH₂CH₃)

Which pair of reagents would produce methyl propanoate (CH₃CH₂COOCH₃) in an esterification?
A. Propanoic acid + methanol
B. Ethanoic acid + methanol
C. Propanoic acid + ethanol
D. Butanoic acid + methanol

A polyester is formed by condensation between a diacid and a diol. Which monomer pair could give poly(ethylene terephthalate) (PET)-type backbone?
A. Ethane-1,2-diol and benzene-1,4-dicarboxylic acid (terephthalic acid)
B. Ethanol and acetic acid
C. Propane-1,2-diol and methanoic acid
D. Ethane and ethene

In acid-catalysed hydrolysis of an ester, which statements are correct?

It is a reversible reaction.

Water acts as a nucleophile.

The reaction rate increases with increasing [H⁺].
A. 1 and 2 only
B. 1, 2 and 3
C. 2 and 3 only
D. 1 and 3 only

Which repeating unit represents the structure of nylon-6,6? (choose description)
A. –NH–(CH₂)₆–NH–CO–(CH₂)₄–CO–
B. –NH–(CH₂)₆–NH–CO–(CH₂)₆–CO–
C. –O–CH₂–CH₂–O–CO–CH₂–CH₂–CO–
D. –NH–CH₂–CH₂–NH–CO–CH₂–CH₂–CO–

Which compound is commonly used to remove SO₂ from flue gases in wet scrubbing processes?
A. Calcium hydroxide slurry (lime)
B. Nitric acid solution
C. Copper(II) sulfate solution
D. Sodium chloride brine
""",

    "Pure Physics": """
Example O-Level Pure Physics MCQs:
I. Measurement

1. A vernier caliper has a least count of 0.01 cm. The main scale reads 2.30 cm and the 7th division of the vernier scale coincides with a main scale mark. What is the correct length?
A. 2.37 cm
B. 2.36 cm
C. 2.307 cm
D. 2.70 cm

2. A micrometer screw gauge has a zero error of –0.02 mm. If it reads 5.46 mm for a wire, what is the true diameter?
A. 5.44 mm
B. 5.46 mm
C. 5.48 mm
D. 5.50 mm

II. Newtonian Mechanics
Kinematics

3. A ball is thrown vertically upward with velocity 20 m/s. After how many seconds will it reach a height of 15 m? (g = 10 m/s²)
A. 1.0 s or 3.0 s
B. 1.5 s only
C. 2.0 s or 2.5 s
D. 2.0 s only

4. A car travels the first 200 m in 20 s at constant acceleration from rest. What is its final velocity after 20 s?
A. 10 m/s
B. 15 m/s
C. 20 m/s
D. 25 m/s

Dynamics

5. A 4.0 kg box is pulled by a force of 30 N at 37° to the horizontal. The horizontal surface has a friction force of 12 N. What is the horizontal acceleration?
A. 2.0 m/s²
B. 3.0 m/s²
C. 4.0 m/s²
D. 5.0 m/s²

6. A ball of mass 0.50 kg moving at 4.0 m/s strikes a wall and rebounds with velocity 3.0 m/s in the opposite direction. What is the impulse on the ball?
A. 0.50 Ns
B. 1.0 Ns
C. 3.5 Ns
D. 0.25 Ns

Turning Effect of Forces

7. A uniform beam of length 4.0 m and weight 200 N is supported at its ends. A 100 N weight is placed 1.0 m from the left support. What is the reaction force at the left support?
A. 100 N
B. 150 N
C. 200 N
D. 250 N

8. A wheel of radius 0.40 m is rotated by a force of 25 N applied tangentially. What is the moment of the force about the wheel's centre?
A. 5.0 Nm
B. 7.5 Nm
C. 10 Nm
D. 12.5 Nm

Pressure

9. A diver is 25 m below the surface of the sea (density of seawater = 1.03 × 10³ kg/m³). What is the pressure due to the water at this depth? (g = 10 m/s²)
A. 2.5 × 10⁴ Pa
B. 2.6 × 10⁵ Pa
C. 2.0 × 10⁵ Pa
D. 2.57 × 10⁵ Pa

10. A piston of area 0.020 m² is used in a hydraulic press to lift a car of weight 10 000 N. If the input force applied is 500 N, what must be the ratio of output piston area to input piston area?
A. 10
B. 20
C. 25
D. 50

Energy

11. A 2.0 kg mass is lifted vertically at a constant speed through 3.0 m in 4.0 s. What is the useful power developed?
A. 6.0 W
B. 10 W
C. 12 W
D. 15 W

12. A motor of efficiency 80% lifts a 500 N load through a height of 12 m in 15 s. What is the input power to the motor?
A. 400 W
B. 500 W
C. 600 W
D. 700 W

III. Thermal Physics
Kinetic Particle Model

13. Which statement best explains why gas pressure decreases when the gas is cooled at constant volume?
A. Particles collide less frequently with the container walls.
B. The number of particles decreases.
C. The average mass of the particles decreases.
D. The container shrinks.

14. Which process requires intermolecular forces to be broken without a change in temperature?
A. Heating water from 20 °C to 80 °C
B. Melting ice at 0 °C
C. Cooling water vapor from 110 °C to 100 °C
D. Cooling liquid nitrogen from –180 °C to –190 °C

Thermal Processes

15. Which method of heat transfer can occur without a medium?
A. Conduction
B. Convection
C. Evaporation
D. Radiation

16. Why is the inside of a vacuum flask silvered?
A. To reduce conduction
B. To reduce convection
C. To reduce radiation
D. To absorb more heat

Thermal Properties

17. 500 J of heat is supplied to 0.10 kg of copper, raising its temperature from 20 °C to 40 °C. What is the specific heat capacity of copper?
A. 100 J/kg°C
B. 200 J/kg°C
C. 250 J/kg°C
D. 500 J/kg°C

18. A block of ice at 0 °C absorbs 3340 J of energy and melts completely. What is its mass? (Lf = 334 J/g)
A. 0.01 kg
B. 0.10 kg
C. 1.0 kg
D. 10 kg

IV. Waves

19. A ripple tank produces water waves of frequency 5.0 Hz and wavelength 0.12 m. What is the speed of the waves?
A. 0.24 m/s
B. 0.40 m/s
C. 0.60 m/s
D. 1.20 m/s

20. A sound wave in air has frequency 340 Hz and wavelength 1.0 m. What is the speed of sound in air from this data?
A. 300 m/s
B. 330 m/s
C. 340 m/s
D. 350 m/s

21. Which property of a sound wave increases when the loudness increases?
A. Amplitude
B. Frequency
C. Wavelength
D. Speed

22. Light of wavelength 5.0 × 10⁻⁷ m enters glass of refractive index 1.5. What is its wavelength in the glass?
A. 1.67 × 10⁻⁷ m
B. 3.33 × 10⁻⁷ m
C. 5.0 × 10⁻⁷ m
D. 7.5 × 10⁻⁷ m

V. Electricity & Magnetism
Static Electricity

23. A plastic rod rubbed with wool becomes negatively charged because…
A. protons are transferred from the wool to the rod.
B. electrons are transferred from the rod to the wool.
C. electrons are transferred from the wool to the rod.
D. protons are transferred from the rod to the wool.

Current of Electricity

24. A current of 3.0 A flows through a lamp for 4.0 minutes. How many coulombs of charge pass through the lamp?
A. 12 C
B. 180 C
C. 720 C
D. 12 000 C

25. A heater has resistance 24 Ω and is connected to a 12 V battery. What is the power dissipated in the heater?
A. 3.0 W
B. 6.0 W
C. 12 W
D. 6.0 J/s

D.C. Circuits

26. A 12 V battery is connected across two resistors in series: 2.0 Ω and 4.0 Ω. What is the current in the circuit?
A. 1.0 A
B. 2.0 A
C. 3.0 A
D. 6.0 A

27. Two identical bulbs are connected in parallel across a 6.0 V supply. Each bulb has resistance 12 Ω. What is the total current supplied by the source?
A. 0.25 A
B. 0.50 A
C. 1.0 A
D. 2.0 A

Practical Electricity

28. An electric kettle is rated 2000 W, 240 V. What is the current in the kettle when operating normally?
A. 4.2 A
B. 6.0 A
C. 8.3 A
D. 10 A

29. A household circuit is protected by a 13 A fuse. Which of these appliances can safely be used on this circuit?
A. An iron rated 3.0 kW at 240 V
B. A kettle rated 2.0 kW at 240 V
C. A heater rated 3.5 kW at 240 V
D. An oven rated 5.0 kW at 240 V

Magnetism

30. Which material is suitable for making the core of a temporary electromagnet?
A. Steel
B. Copper
C. Soft iron
D. Aluminium

Electromagnetism

31. The direction of force on a current-carrying conductor in a magnetic field is given by…
A. Fleming's left-hand rule.
B. Fleming's right-hand rule.
C. Newton's third law.
D. Lenz's law.

Electromagnetic Induction

32. A coil of 200 turns is rotated in a magnetic field at double the speed. How does the induced e.m.f. change?
A. Halved
B. Doubled
C. Same
D. Zero

33. A transformer has 500 turns on the primary coil and 50 turns on the secondary. If the primary voltage is 240 V, what is the secondary voltage?
A. 12 V
B. 24 V
C. 48 V
D. 2400 V

VI. Radioactivity

34. Which radiation has the greatest penetrating power?
A. α
B. β
C. γ
D. All equal

35. A sample of radioactive material has a half-life of 20 minutes. What fraction of the original sample remains after 1 hour?
A. 1/2
B. 1/3
C. 1/4
D. 1/8

36. The unit becquerel (Bq) is defined as…
A. one emission per second.
B. one joule per second.
C. one coulomb per second.
D. one proton per second.

37. A radioactive isotope emits only β-particles. What change occurs in the nucleus?
A. The proton number decreases by 1.
B. The proton number increases by 1.
C. The nucleon number decreases by 1.
D. The nucleon number increases by 1.

38. Which radiation is used in smoke detectors?
A. α
B. β
C. γ
D. X-rays

39. Which statement about background radiation is correct?
A. It comes only from cosmic rays.
B. It is due only to human activities.
C. It includes natural sources such as rocks and soil.
D. It cannot be measured.

40. Which radiation is deflected the least in an electric field?
A. α-particles
B. β-particles
C. γ-rays
D. All are equally deflected
""",

    "Pure Biology": """
Example O-Level Pure Biology MCQs:
1. Which structure controls the passage of substances in and out of the cell? 
A. Nucleus 
B. Cytoplasm 
C. Cell membrane 
D. Mitochondrion 
...
""",

    "EMath": "Example O-Level EMath MCQs:\n1. [你的例题在这里]\n...",
    "AMath": "Example O-Level AMath MCQs:\n1. [你的例题在这里]\n...",
    "Combine Chemistry": "Example O-Level Combine Chemistry MCQs:\n1. [你的例题在这里]\n...",
    "Combine Physics": "Example O-Level Combine Physics MCQs:\n1. [你的例题在这里]\n...",
    "Combine Biology": "Example O-Level Combine Biology MCQs:\n1. [你的例题在这里]\n...",
    "Combine Chemistry and Biology": """Example O-Level Combine Chemistry and Biology MCQs (half chemistry, half biology):
Chemistry (20 Questions)
Matter – Structures and Properties

1. Which statement best explains why graphite conducts electricity but diamond does not?
A. Graphite has delocalised electrons that can move between layers.
B. Graphite atoms are more electronegative than diamond atoms.
C. Diamond has free-moving ions in its structure.
D. Diamond has weaker covalent bonds than graphite.

2. Which property of giant ionic lattices decreases as the size of the cation increases?
A. Lattice energy
B. Melting point
C. Strength of electrostatic attraction
D. Solubility in water

Chemical Bonding and Structure

3. Which molecule contains both covalent and dative covalent (coordinate) bonds?
A. CO₂
B. NH₄⁺
C. CH₄
D. HCl

4. Which type of intermolecular forces is strongest in hydrogen fluoride, HF?
A. Permanent dipole-dipole
B. Hydrogen bonding
C. Van der Waals forces
D. Ionic bonding

Chemical Calculations

5. A compound contains 27.3% carbon, 72.7% oxygen by mass. What is its empirical formula?
A. CO
B. CO₂
C. C₂O₃
D. C₂O₄

6. 0.500 mol of a gas occupies 12.0 dm³ at room temperature and pressure. What volume will 1.00 mol occupy under the same conditions?
A. 6.0 dm³
B. 12.0 dm³
C. 24.0 dm³
D. 48.0 dm³

Acid-Base Chemistry

7. Which indicator is most suitable for titrating a weak acid with a strong alkali?
A. Methyl orange
B. Phenolphthalein
C. Universal indicator
D. Litmus

8. Which solution contains the same concentration of hydroxide ions as 0.01 mol dm⁻³ NaOH?
A. 0.005 mol dm⁻³ Ca(OH)₂
B. 0.01 mol dm⁻³ HCl
C. 0.02 mol dm⁻³ NH₃
D. 0.01 mol dm⁻³ CH₃COOH

Qualitative Analysis

9. Which observation identifies the presence of Fe²⁺ ions?
A. Green precipitate with NaOH that turns brown on standing
B. Blue precipitate with NaOH that dissolves in excess
C. White precipitate with NaOH that dissolves in excess
D. Brown gas released when heated with HCl

10. Which gas would turn acidified potassium dichromate(VI) solution from orange to green?
A. Ammonia
B. Sulfur dioxide
C. Carbon dioxide
D. Oxygen

Redox Chemistry

11. In the reaction:
Zn(s) + Cu²⁺(aq) → Zn²⁺(aq) + Cu(s),
which statement is correct?
A. Zinc is oxidised and acts as the reducing agent.
B. Zinc is reduced and acts as the oxidising agent.
C. Copper is oxidised and acts as the reducing agent.
D. Copper is reduced and acts as the reducing agent.

Periodic Table

12. Which oxide is amphoteric?
A. SiO₂
B. Na₂O
C. Al₂O₃
D. P₂O₅

13. Which halogen reacts most vigorously with iron to form iron(III) halide?
A. Fluorine
B. Chlorine
C. Bromine
D. Iodine

Energetics

14. Which process is exothermic?
A. Melting ice
B. Combustion of ethanol
C. Electrolysis of water
D. Thermal decomposition of CaCO₃

15. Which energy change corresponds to the enthalpy of neutralisation?
A. Heat absorbed when one mole of water is formed from neutralisation
B. Heat released when one mole of an alkali is dissolved in water
C. Heat released when one mole of a salt is dissolved in water
D. Heat absorbed when acid is titrated with alkali

Rates of Reaction

16. Which change increases the rate of reaction between magnesium and dilute HCl but does not affect the total volume of hydrogen produced?
A. Increasing temperature
B. Increasing concentration of HCl
C. Using powdered magnesium instead of ribbon
D. All of the above

17. Which reaction would be slowed most if the reaction vessel is cooled?
A. Thermal decomposition of CaCO₃
B. Neutralisation of HCl with NaOH
C. Combustion of methane
D. Reaction between Zn and H₂SO₄

Organic Chemistry

18. Which compound does not decolourise bromine water in the dark?
A. Ethene
B. Propene
C. Ethane
D. But-2-ene

19. Which statement about alcohols is correct?
A. They undergo substitution with sodium to form hydrogen gas.
B. They are weaker acids than carboxylic acids.
C. They cannot form hydrogen bonds with water.
D. They react with carbonates to produce CO₂.

20. Which type of polymer is formed when amino acids join together?
A. Addition polymer
B. Condensation polymer
C. Polyester
D. Polyalkene

Biology (20 Questions)
Cells & Life Chemistry

21. Which organelle is responsible for the modification and packaging of proteins?
A. Nucleus
B. Golgi apparatus
C. Mitochondrion
D. Endoplasmic reticulum

22. Which statement about diffusion is correct?
A. It requires energy in the form of ATP.
B. It is faster at lower temperatures.
C. It moves substances against the concentration gradient.
D. It occurs because of random movement of particles.

23. Which biological molecule is correctly matched to its subunit?
A. Protein – glucose
B. Lipid – fatty acids and glycerol
C. Starch – amino acids
D. DNA – fatty acids

Human Nutrition & Transport

24. Which enzyme digests proteins into peptides in the stomach?
A. Amylase
B. Lipase
C. Pepsin
D. Trypsin

25. Which component of blood is responsible for transporting carbon dioxide mainly as hydrogen carbonate ions?
A. Plasma
B. Platelets
C. Red blood cells
D. White blood cells

26. Which chamber of the heart pumps blood at the highest pressure?
A. Left atrium
B. Left ventricle
C. Right atrium
D. Right ventricle

Respiration

27. Which gas exchange adaptation is found in both alveoli and root hair cells?
A. Large surface area
B. Presence of haemoglobin
C. Thin layer of moisture
D. Active transport

28. Which equation shows anaerobic respiration in yeast?
A. Glucose → lactic acid + energy
B. Glucose → ethanol + carbon dioxide + energy
C. Glucose → carbon dioxide + water + energy
D. Glucose → glucose-6-phosphate

Infectious Diseases

29. Which disease is prevented by the BCG vaccine?
A. Malaria
B. Tuberculosis
C. Cholera
D. Polio

30. Which process occurs during phagocytosis?
A. Production of antibodies
B. Engulfing of pathogens
C. Neutralisation of toxins
D. Memory cell formation

Plants & Ecosystems

31. Which tissue in plants transports mainly sucrose and amino acids?
A. Xylem
B. Phloem
C. Cambium
D. Cortex

32. Which statement explains why plants in shaded forests often have broad leaves?
A. To increase transpiration rate
B. To reduce water loss
C. To maximise light absorption
D. To prevent herbivory

33. Which process is represented by:
ammonium ions → nitrite ions → nitrate ions?
A. Nitrogen fixation
B. Nitrification
C. Denitrification
D. Ammonification

34. Which organisms recycle carbon by converting dead organic matter into carbon dioxide?
A. Decomposers
B. Producers
C. Herbivores
D. Carnivores

Genetics & Reproduction

35. Which statement about meiosis is correct?
A. It produces genetically identical daughter cells.
B. It results in cells with the same chromosome number as the parent cell.
C. It produces gametes with half the chromosome number.
D. It occurs in all body cells.

36. In DNA, the base adenine pairs with:
A. cytosine
B. guanine
C. thymine
D. uracil

37. Which process ensures that offspring show variation?
A. Mitosis
B. Fertilisation
C. DNA replication
D. Cytokinesis

38. Which hormone stimulates ovulation?
A. Oestrogen
B. Progesterone
C. Follicle-stimulating hormone (FSH)
D. Luteinising hormone (LH)

39. Which type of reproduction produces offspring genetically identical to the parent?
A. Sexual reproduction
B. Asexual reproduction
C. Fertilisation
D. Pollination

40. A man with blood group AB has a child with a woman who has blood group O. What possible blood groups can their child have?
A. A and B only
B. A, B, AB, O
C. A, B, O only
D. AB only
""",
}

def generate_mcq_questions(subject="Pure Chemistry", num_questions=10, difficulty="same"):
    """生成MCQ问题"""
    examples = SUBJECT_EXAMPLES.get(subject, "Example O-Level MCQs:\n")
    prompt = f"""{examples}

Now, generate {num_questions} NEW O-Level {subject} multiple choice questions
in the same style and {difficulty} difficulty. 
Each question must have 4 options (A–D).
Do not repeat the examples.

Instructions:
1. Do NOT repeat the examples or the specific topics they contain.
2. Replace specific substances, elements, or chemical names with generic labels (X, Y, Z) or alternative examples.
3. Use the same question types (calculation, reasoning, concept-based) and format.
4. Make sure the new questions are original, varied, and cover general concepts only.
"""
    
    print(f"🚀 Starting MCQ generation for {subject}...")
    print(f"📊 Generating {num_questions} questions with {difficulty} difficulty")
    print("⏳ Please wait while connecting to AI API...")
    
    start_time = time.time()
    
    try:
        response = client.chat.completions.create(
            model="mistralai/mistral-7b-instruct:free",
            messages=[{"role": "user", "content": prompt}]
        )
        
        generation_time = time.time() - start_time
        print(f"✅ MCQ generation completed successfully! Time taken: {generation_time:.2f} seconds")
        
        return response.choices[0].message.content
    except Exception as e:
        print(f"❌ Error during MCQ generation: {e}")
        return None


def generate_mcq_answers(questions):
    """生成MCQ答案"""
    if not questions:
        print("❌ No questions provided for answer generation")
        return None
        
    prompt = f"""
Provide only the correct answers with explanation for these questions:

{questions}

Instructions:
1. Format each answer as: QuestionNumber. CorrectOption — Explanation
   Example: 1. B — The calculation shows that...
2. Keep explanations concise (1–2 sentences each).
3. Avoid repeating specific chemicals, elements, or compounds from the original questions; use generic labels if necessary.
4. Separate each answer by a blank line.
5. Use clear and simple English suitable for O-Level students.
6. Output only the text, do not include extra commentary or notes.
"""
    
    print("🧠 Starting answer generation...")
    print("⏳ Generating answers and explanations...")
    
    start_time = time.time()
    
    try:
        response = client.chat.completions.create(
            model="mistralai/mistral-7b-instruct:free",
            messages=[{"role": "user", "content": prompt}]
        )
        
        generation_time = time.time() - start_time
        print(f"✅ Answer generation completed! Time taken: {generation_time:.2f} seconds")
        
        return response.choices[0].message.content
    except Exception as e:
        print(f"❌ Error during answer generation: {e}")
        return None


def save_to_word(text, filename):
    """保存文本到Word文档"""
    if not text:
        print(f"⚠️  No content to save for {filename}")
        return False
        
    try:
        doc = Document()
        # 添加标题
        title = os.path.basename(filename).replace('.docx', '')
        doc.add_heading(title, level=1)
        
        # 添加内容
        for line in text.split("\n"):
            if line.strip():  # 跳过空行
                doc.add_paragraph(line)
        
        doc.save(filename)
        print(f"💾 File saved successfully: {filename}")
        return True
    except Exception as e:
        print(f"❌ Error saving file {filename}: {e}")
        return False


def create_folder_with_mcq(folder_name, subject="Pure Chemistry", num_questions=10, difficulty="same"):
    """创建文件夹并生成MCQ"""
    print("=" * 60)
    print(f"🎯 STARTING MCQ GENERATION PROCESS")
    print("=" * 60)
    print(f"📁 Folder: {folder_name}")
    print(f"📚 Subject: {subject}")
    print(f"🔢 Number of questions: {num_questions}")
    print(f"🎚️  Difficulty: {difficulty}")
    print("-" * 60)
    
    # 创建文件夹
    try:
        os.makedirs(folder_name, exist_ok=True)
        print(f"📂 Folder created/verified: {os.path.abspath(folder_name)}")
    except Exception as e:
        print(f"❌ Error creating folder: {e}")
        return

    # 生成问题
    questions = generate_mcq_questions(subject, num_questions, difficulty)
    
    if questions:
        # 保存问题文件
        questions_file = os.path.join(folder_name, f"{subject}_MCQ.docx")
        questions_saved = save_to_word(questions, questions_file)
        
        if questions_saved:
            # 生成答案
            answers = generate_mcq_answers(questions)
            
            if answers:
                # 保存答案文件
                answers_file = os.path.join(folder_name, f"{subject}_Answers.docx")
                answers_saved = save_to_word(answers, answers_file)
                
                if answers_saved:
                    print("=" * 60)
                    print("🎉 PROCESS COMPLETED SUCCESSFULLY!")
                    print("=" * 60)
                    print(f"📋 Generated files:")
                    print(f"   • Questions: {questions_file}")
                    print(f"   • Answers: {answers_file}")
                    print(f"📁 Full path: {os.path.abspath(folder_name)}")
                    print("=" * 60)
                else:
                    print("❌ Failed to save answers file")
            else:
                print("❌ Failed to generate answers")
        else:
            print("❌ Failed to save questions file")
    else:
        print("❌ Failed to generate questions")
    
    print("🛑 Process stopped.")


# 示例运行
if __name__ == "__main__":
    print("🔧 MCQ Generator Tool - O Level Question Bank")
    print("Initializing system...")
    
    # 示例运行
    create_folder_with_mcq("OLevel_CombineScience", subject="Combine Chemistry and Biology", num_questions=40, difficulty="much higher")